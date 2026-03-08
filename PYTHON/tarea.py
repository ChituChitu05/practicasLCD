# Crear documento Word con los puntos solicitados

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Crear documento
doc = Document()
doc.add_heading('Actividad 3 - Definiciones de Economía', level=1)

# 3a
doc.add_heading('3a. Cinco definiciones de Economía', level=2)

table = doc.add_table(rows=6, cols=2)
table.rows[0].cells[0].text = "Autor"
table.rows[0].cells[1].text = "Definición"

definitions = [
    ("Lionel Robbins (1932)", 
     "La economía es la ciencia que estudia la conducta humana como una relación entre fines y medios escasos que tienen usos alternativos."),
    ("Paul A. Samuelson (1948)", 
     "La economía es el estudio de cómo las sociedades utilizan recursos escasos para producir bienes valiosos y distribuirlos entre diferentes personas."),
    ("Gregory Mankiw (1998)", 
     "La economía es el estudio de cómo la sociedad administra sus recursos escasos."),
    ("Alfred Marshall (1890)", 
     "La Economía Política o Economía es el estudio de la humanidad en los asuntos ordinarios de la vida."),
    ("Adam Smith (1776)", 
     "Una investigación sobre la naturaleza y las causas de la riqueza de las naciones.")
]

for i, (author, definition) in enumerate(definitions, start=1):
    table.rows[i].cells[0].text = author
    table.rows[i].cells[1].text = definition

# 3b
doc.add_heading('3b. Orden de las definiciones (de más completa a menos completa)', level=2)

order = [
    "1. Paul A. Samuelson",
    "2. Lionel Robbins",
    "3. Gregory Mankiw",
    "4. Alfred Marshall",
    "5. Adam Smith"
]

for item in order:
    doc.add_paragraph(item)

# 3c
doc.add_heading('3c. Palabras clave (subrayadas con marcatexto)', level=2)

keywords = [
    "Fines", "Medios escasos", "Usos alternativos",
    "Recursos escasos", "Producir bienes", "Distribuir",
    "Administrar", "Asuntos ordinarios de la vida", "Riqueza"
]

p = doc.add_paragraph("Palabras clave identificadas: ")

for word in keywords:
    run = doc.add_paragraph().add_run(word)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# Guardar documento
file_path = "/mnt/data/Actividad_3_Definiciones_de_Economia.docx"
doc.save(file_path)

file_path

